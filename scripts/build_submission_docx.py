"""Build the LMS submission DOCX from the current Markdown project brief."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submission" / "PROJECT_BRIEF.md"
OUTPUT = ROOT / "submission" / "Submission_Dilnura_Hamdamova.docx"


def clean_markup(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def build() -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    in_code = False
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line:
            document.add_paragraph()
        elif in_code:
            paragraph = document.add_paragraph(style="No Spacing")
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif line.startswith("# "):
            document.add_heading(clean_markup(line[2:]), level=0)
        elif line.startswith("## "):
            document.add_heading(clean_markup(line[3:]), level=1)
        elif line.startswith("### "):
            document.add_heading(clean_markup(line[4:]), level=2)
        elif line.startswith("- "):
            document.add_paragraph(clean_markup(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            document.add_paragraph(clean_markup(line.split(". ", 1)[1]), style="List Number")
        else:
            document.add_paragraph(clean_markup(line))

    document.core_properties.title = "Tashkent Apartment Listing Price Predictor — Project Brief"
    document.core_properties.author = "Dilnura Hamdamova"
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
