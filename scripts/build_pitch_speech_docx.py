"""Build the printable pitch speech DOCX from the committed Markdown script."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "defense_pitch_script_en.md"
OUTPUT = ROOT / "presentation" / "Tashkent_Apartment_Pitch_Speech.docx"


def add_inline_runs(paragraph, text: str) -> None:
    """Render the small subset of inline Markdown used by the speech."""
    parts = text.split("**")
    for index, part in enumerate(parts):
        run = paragraph.add_run(part.replace("`", ""))
        run.bold = index % 2 == 1


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    code_mode = False
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            code_mode = not code_mode
            continue
        if line == "<!-- PAGEBREAK -->":
            document.add_page_break()
            continue
        if not line:
            continue
        if line.startswith("# "):
            paragraph = document.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            heading = document.add_heading(line[3:], level=1)
            heading.paragraph_format.keep_with_next = True
        elif line.startswith("> "):
            paragraph = document.add_paragraph(style="Quote")
            add_inline_runs(paragraph, line[2:])
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, line[2:])
        elif code_mode:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, line)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Dilnura Hamdamova — 2026 Tashkent Apartment Price Predictor")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
